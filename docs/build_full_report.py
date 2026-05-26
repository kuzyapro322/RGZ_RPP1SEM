from pathlib import Path
from docx import Document
from docx.shared import Inches, Pt
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.enum.section import WD_SECTION_START
from docx.oxml import OxmlElement
from docx.oxml.ns import qn

ROOT = Path.cwd()
DOCS = ROOT / 'docs'
SCREENSHOTS = DOCS / 'screenshots'
LOGO = ROOT / 'web' / 'static' / 'img' / 'nstu_logo.png'
OUT = DOCS / 'RGZ_variant_7_full_report.docx'


def set_cell_text(cell, text, bold=False):
    cell.text = ''
    p = cell.paragraphs[0]
    r = p.add_run(text)
    r.bold = bold
    for run in p.runs:
        run.font.name = 'Times New Roman'
        run.font.size = Pt(12)


def set_margins(section):
    section.top_margin = Inches(0.8)
    section.bottom_margin = Inches(0.8)
    section.left_margin = Inches(1.0)
    section.right_margin = Inches(0.7)


def add_page_number(section):
    footer = section.footer
    p = footer.paragraphs[0]
    p.alignment = WD_ALIGN_PARAGRAPH.RIGHT
    run = p.add_run()
    fld_char1 = OxmlElement('w:fldChar')
    fld_char1.set(qn('w:fldCharType'), 'begin')
    instr_text = OxmlElement('w:instrText')
    instr_text.set(qn('xml:space'), 'preserve')
    instr_text.text = 'PAGE'
    fld_char2 = OxmlElement('w:fldChar')
    fld_char2.set(qn('w:fldCharType'), 'end')
    run._r.append(fld_char1)
    run._r.append(instr_text)
    run._r.append(fld_char2)


def add_heading(doc, text, level=1):
    p = doc.add_paragraph()
    if level == 1:
        p.alignment = WD_ALIGN_PARAGRAPH.CENTER
        size = 16
    else:
        p.alignment = WD_ALIGN_PARAGRAPH.LEFT
        size = 14
    r = p.add_run(text)
    r.bold = True
    r.font.name = 'Times New Roman'
    r.font.size = Pt(size)
    return p


def add_paragraph(doc, text):
    p = doc.add_paragraph()
    p.alignment = WD_ALIGN_PARAGRAPH.JUSTIFY
    p.paragraph_format.first_line_indent = Inches(0.49)
    p.paragraph_format.line_spacing = 1.5
    p.paragraph_format.space_after = Pt(0)
    p.paragraph_format.space_before = Pt(0)
    r = p.add_run(text)
    r.font.name = 'Times New Roman'
    r.font.size = Pt(14)
    return p


def add_caption(doc, text):
    p = doc.add_paragraph()
    p.alignment = WD_ALIGN_PARAGRAPH.CENTER
    p.paragraph_format.space_before = Pt(6)
    p.paragraph_format.space_after = Pt(6)
    r = p.add_run(text)
    r.font.name = 'Times New Roman'
    r.font.size = Pt(12)
    r.italic = True
    return p


def add_image(doc, filename, width=6.3):
    p = doc.add_paragraph()
    p.alignment = WD_ALIGN_PARAGRAPH.CENTER
    run = p.add_run()
    run.add_picture(str(SCREENSHOTS / filename), width=Inches(width))
    return p


DOCS.mkdir(exist_ok=True)
doc = Document()
set_margins(doc.sections[0])
add_page_number(doc.sections[0])

styles = doc.styles
styles['Normal'].font.name = 'Times New Roman'
styles['Normal'].font.size = Pt(14)

# Титульный лист
p = doc.add_paragraph()
p.alignment = WD_ALIGN_PARAGRAPH.CENTER
p.add_run().add_picture(str(LOGO), width=Inches(3.0))

for line, size, bold in [
    ('Расчетно-графическое задание', 16, True),
    ('по дисциплине «Разработка программных приложений»', 14, False),
    ('Тема: «Разработка WEB-приложения для учета финансов»', 14, False),
    ('Индивидуальный вариант: 7 — удаление операции', 14, True),
]:
    p = doc.add_paragraph()
    p.alignment = WD_ALIGN_PARAGRAPH.CENTER
    r = p.add_run(line)
    r.font.name = 'Times New Roman'
    r.font.size = Pt(size)
    r.bold = bold

p = doc.add_paragraph()
p.alignment = WD_ALIGN_PARAGRAPH.RIGHT
r = p.add_run('Выполнили: студенты группы ________\nПроверил: __________________________')
r.font.name = 'Times New Roman'
r.font.size = Pt(14)

p = doc.add_paragraph()
p.alignment = WD_ALIGN_PARAGRAPH.CENTER
p.paragraph_format.space_before = Pt(120)
r = p.add_run('Новосибирск, 2026')
r.font.name = 'Times New Roman'
r.font.size = Pt(14)

doc.add_page_break()

add_heading(doc, 'Содержание', 1)
for item in [
    'Введение',
    '1 Теоретические основы разработки WEB-приложения',
    '1.1 WEB-приложение и клиент-серверная архитектура',
    '1.2 Flask и маршрутизация HTTP-запросов',
    '1.3 Jinja-шаблоны и серверная генерация страниц',
    '1.4 PostgreSQL и SQLAlchemy ORM',
    '1.5 Хэширование паролей и сессии пользователя',
    '1.6 Docker Compose и развертывание сервисов',
    '2 Практическая реализация приложения',
    '2.1 Структура проекта',
    '2.2 Схема базы данных',
    '2.3 Реализация общей части задания',
    '2.4 Реализация индивидуального варианта 7',
    '2.5 Демонстрация работы приложения',
    'Заключение',
    'Список использованных источников',
]:
    add_paragraph(doc, item)

doc.add_page_break()

add_heading(doc, 'Введение', 1)
add_paragraph(doc, 'Целью расчетно-графического задания является разработка WEB-приложения для учета личных финансов. Приложение должно позволять пользователю зарегистрироваться, добавлять доходные и расходные операции, просматривать список операций и получать суммы в выбранной валюте. В качестве базы данных используется PostgreSQL, а пользовательский интерфейс формируется с помощью Jinja-шаблонов.')
add_paragraph(doc, 'В рамках индивидуального варианта 7 дополнительно реализована функция удаления операции. Эта функция важна для финансового приложения, потому что пользователь может ошибиться при вводе суммы, даты или типа операции. Удаление выполняется только для операций текущего авторизованного пользователя, что защищает данные других пользователей.')
add_paragraph(doc, 'Практическая часть выполнена в виде нескольких Docker-сервисов: основного WEB-приложения, базы данных PostgreSQL и внешнего сервиса курса валют. Такой подход позволяет запускать проект одной командой и демонстрировать работу приложения в условиях, близких к реальному развертыванию [5].')

add_heading(doc, '1 Теоретические основы разработки WEB-приложения', 1)
add_heading(doc, '1.1 WEB-приложение и клиент-серверная архитектура', 2)
add_paragraph(doc, 'WEB-приложение работает по клиент-серверной модели: пользователь открывает страницу в браузере, браузер отправляет HTTP-запрос, сервер обрабатывает его и возвращает HTTP-ответ. В ответе может находиться HTML-страница, JSON-данные или перенаправление на другую страницу. В данной работе браузер является клиентом, а Flask-приложение выполняет роль сервера.')
add_paragraph(doc, 'Для учета финансов выбран классический серверный подход: HTML-страницы формируются на backend-стороне, а затем отправляются пользователю. Такой вариант удобен для учебной работы, потому что вся основная логика находится в одном приложении: маршруты, проверка авторизации, работа с базой данных и отображение результатов.')

add_heading(doc, '1.2 Flask и маршрутизация HTTP-запросов', 2)
add_paragraph(doc, 'Flask — это Python-фреймворк для создания WEB-приложений. Он позволяет описывать маршруты, которые связывают URL-адрес и Python-функцию. Например, маршрут /reg отвечает за регистрацию пользователя, /add_operation — за добавление операции, а /operations — за отображение списка операций [1].')
add_paragraph(doc, 'В Flask обработчик маршрута получает данные запроса через объект request, может работать с пользовательской сессией через session и возвращать шаблон, JSON-ответ или перенаправление. В приложении используются GET-запросы для отображения страниц и POST-запросы для действий, которые изменяют данные: регистрация, добавление операции, выход из аккаунта и удаление операции.')

add_heading(doc, '1.3 Jinja-шаблоны и серверная генерация страниц', 2)
add_paragraph(doc, 'Jinja — это шаблонизатор, который позволяет формировать HTML-страницы на основе шаблонов и данных, полученных от backend. В шаблонах можно использовать переменные, циклы и условия. Например, на странице операций через цикл выводятся все операции пользователя, а через условие отображается сообщение, если операций еще нет [2].')
add_paragraph(doc, 'Использование Jinja удобно для данного задания, потому что frontend должен быть реализован именно через шаблоны. В проекте создан базовый шаблон base.html, от которого наследуются страницы регистрации, авторизации, добавления операции и просмотра операций. Это уменьшает повторение HTML-кода и делает интерфейс единообразным.')

add_heading(doc, '1.4 PostgreSQL и SQLAlchemy ORM', 2)
add_paragraph(doc, 'PostgreSQL — это реляционная система управления базами данных. Данные в ней хранятся в таблицах, а связи между сущностями задаются с помощью первичных и внешних ключей. В данной работе используются таблицы users и operations: первая хранит пользователей, вторая — финансовые операции [4].')
add_paragraph(doc, 'Для работы с базой данных используется SQLAlchemy ORM. ORM позволяет описывать таблицы как Python-классы, а записи таблиц — как объекты. Благодаря этому в коде можно создавать пользователя или операцию через классы User и Operation, а SQLAlchemy самостоятельно формирует SQL-запросы к PostgreSQL [3].')

add_heading(doc, '1.5 Хэширование паролей и сессии пользователя', 2)
add_paragraph(doc, 'Пароль нельзя хранить в базе данных в открытом виде. При регистрации приложение создает хэш пароля с помощью generate_password_hash, а при входе проверяет пароль через check_password_hash. Такой подход снижает риск раскрытия паролей при несанкционированном доступе к базе данных [6].')
add_paragraph(doc, 'После успешной регистрации или авторизации идентификатор пользователя сохраняется в session. Сессия позволяет серверу понимать, какой пользователь выполняет запрос. Для защищенных страниц используется проверка авторизации: если пользователь не вошел в аккаунт, он перенаправляется на страницу входа.')

add_heading(doc, '1.6 Docker Compose и развертывание сервисов', 2)
add_paragraph(doc, 'Docker Compose используется для запуска нескольких связанных сервисов одной командой. В проекте описаны три сервиса: db для PostgreSQL, web для основного Flask-приложения и rate_service для внешнего сервиса курса валют. Это упрощает запуск и проверку проекта, потому что не нужно вручную настраивать базу данных и окружение Python [5].')
add_paragraph(doc, 'Основное приложение обращается к базе данных по внутреннему имени сервиса db, а к сервису валют — по имени rate_service. Внешний сервис принимает запрос /rate с параметром currency и возвращает курс в JSON-формате. Для выполнения HTTP-запроса из основного приложения используется библиотека requests [7].')

add_heading(doc, '2 Практическая реализация приложения', 1)
add_heading(doc, '2.1 Структура проекта', 2)
add_paragraph(doc, 'Проект расположен в папке RGZ. Основной backend находится в файле web/app.py. Jinja-шаблоны расположены в папке web/templates, стили интерфейса — в web/static/css/style.css, логотип — в web/static/img/nstu_logo.png. Внешний сервис курса валют находится в rate_service/rate_app.py. Настройки Docker Compose описаны в docker-compose.yml.')
add_paragraph(doc, 'В начале Python-файлов добавлены блоки с пояснением импортов, а основные части кода разделены комментариями по пунктам задания. Это сделано для того, чтобы при защите было проще показать, где реализована регистрация, где добавление операции, где просмотр операций и где индивидуальный вариант 7.')

add_heading(doc, '2.2 Схема базы данных', 2)
add_paragraph(doc, 'База данных содержит таблицу users с полями id, name и password_hash. Поле id является первичным ключом, name хранит логин пользователя, password_hash хранит хэш пароля. Логин сделан уникальным, чтобы два пользователя не могли зарегистрироваться под одним именем.')
add_paragraph(doc, 'Таблица operations содержит поля id, date, sum, user_id и type_operation. Поле user_id является внешним ключом на таблицу users. Благодаря этому каждая операция принадлежит конкретному пользователю. При удалении пользователя связанные операции также могут быть удалены через связь ORM.')

table = doc.add_table(rows=1, cols=4)
table.style = 'Table Grid'
hdr = table.rows[0].cells
for i, text in enumerate(['Таблица', 'Поле', 'Тип данных', 'Назначение']):
    set_cell_text(hdr[i], text, True)
rows = [
    ('users', 'id', 'Integer', 'Уникальный идентификатор пользователя'),
    ('users', 'name', 'String', 'Логин пользователя'),
    ('users', 'password_hash', 'String', 'Хэш пароля'),
    ('operations', 'id', 'Integer', 'Уникальный идентификатор операции'),
    ('operations', 'date', 'Date', 'Дата операции'),
    ('operations', 'sum', 'Numeric(12, 2)', 'Сумма операции в рублях'),
    ('operations', 'user_id', 'Integer', 'Связь с пользователем'),
    ('operations', 'type_operation', 'String', 'Тип операции: income или expense'),
]
for row in rows:
    cells = table.add_row().cells
    for i, text in enumerate(row):
        set_cell_text(cells[i], text)

add_heading(doc, '2.3 Реализация общей части задания', 2)
add_paragraph(doc, 'Регистрация реализована по адресу /reg. При GET-запросе отображается HTML-форма регистрации, при POST-запросе backend получает логин и пароль, проверяет заполненность полей, проверяет отсутствие пользователя с таким логином и сохраняет нового пользователя в PostgreSQL. Если запрос отправлен в JSON-формате, приложение возвращает JSON-ответ со статусом создания пользователя.')
add_paragraph(doc, 'Добавление операции реализовано по адресу /add_operation. Пользователь выбирает тип операции, вводит сумму и дату. Backend проверяет, что пользователь авторизован, что тип операции является income или expense, что сумма больше нуля, после чего сохраняет операцию в таблицу operations.')
add_paragraph(doc, 'Просмотр операций реализован по адресу /operations. Пользователь может выбрать валюту RUB, USD или EUR. Если выбрана иностранная валюта, backend отправляет GET-запрос во внешний сервис курса валют, получает значение rate и пересчитывает сумму операции из рублей в выбранную валюту. Результаты отображаются в таблице через Jinja-шаблон.')

add_heading(doc, '2.4 Реализация индивидуального варианта 7', 2)
add_paragraph(doc, 'По индивидуальному варианту 7 необходимо реализовать удаление операции пользователя. В приложении это выполнено маршрутом POST /delete_operation/<operation_id>. На странице операций напротив каждой строки находится кнопка «Удалить». При нажатии браузер отправляет POST-запрос с идентификатором операции.')
add_paragraph(doc, 'Перед удалением backend выполняет две проверки. Первая проверка — пользователь должен быть авторизован. Вторая проверка — операция с указанным id должна существовать и принадлежать текущему пользователю. Если операция не найдена или принадлежит другому пользователю, сервер возвращает ошибку. Если проверка успешна, запись удаляется из PostgreSQL, а пользователь перенаправляется обратно на страницу операций с сообщением об успешном удалении.')

add_heading(doc, '2.5 Демонстрация работы приложения', 2)
figs = [
    ('01_registration.png', 'Рисунок 1 — страница регистрации пользователя', 'На рисунке 1 показана форма регистрации. Пользователь вводит логин и пароль. После отправки формы backend создает запись в таблице users, а пароль сохраняется в виде хэша. Это соответствует пункту 2.1.2 задания.'),
    ('02_empty_operations.png', 'Рисунок 2 — список операций после регистрации', 'На рисунке 2 показана страница операций нового пользователя. Так как пользователь только зарегистрирован, список операций пуст. Интерфейс предлагает добавить первую операцию.'),
    ('03_add_operation_form.png', 'Рисунок 3 — форма добавления операции', 'На рисунке 3 показана форма добавления операции. Пользователь выбирает тип операции, вводит сумму в рублях и дату. Данные отправляются на маршрут /add_operation методом POST.'),
    ('04_operation_added.png', 'Рисунок 4 — успешное добавление операции', 'На рисунке 4 показано сообщение об успешном сохранении операции. Это подтверждает, что backend принял данные, прошел проверки и добавил запись в таблицу operations.'),
    ('05_operations_usd.png', 'Рисунок 5 — просмотр операций с конвертацией в USD', 'На рисунке 5 показан список операций пользователя. Выбрана валюта USD, поэтому приложение обратилось к внешнему сервису курса валют и вывело суммы как в рублях, так и в долларах.'),
    ('06_after_delete.png', 'Рисунок 6 — результат удаления операции', 'На рисунке 6 показан результат выполнения индивидуального варианта 7. После нажатия кнопки «Удалить» операция удалена из базы данных, а приложение вывело сообщение об успешном удалении.'),
]
for filename, caption, desc in figs:
    add_image(doc, filename)
    add_caption(doc, caption)
    add_paragraph(doc, desc)

add_heading(doc, 'Заключение', 1)
add_paragraph(doc, 'В результате выполнения расчетно-графического задания было разработано WEB-приложение для учета финансов. Приложение поддерживает регистрацию пользователя, авторизацию, добавление доходных и расходных операций, просмотр операций в разных валютах и обращение к внешнему сервису курса валют.')
add_paragraph(doc, 'Индивидуальный вариант 7 реализован полностью: пользователь может удалить выбранную операцию из списка. При этом backend проверяет авторизацию и принадлежность операции текущему пользователю, что делает удаление корректным и безопасным с точки зрения доступа к данным.')
add_paragraph(doc, 'Проект запускается через Docker Compose, что упрощает демонстрацию работы приложения. В качестве базы данных используется PostgreSQL, frontend реализован через Jinja-шаблоны, а backend написан на Flask. Полученное приложение соответствует основным требованиям задания и может быть использовано для защиты РГЗ.')

add_heading(doc, 'Список использованных источников', 1)
sources = [
    '[1] Flask Documentation. URL: https://flask.palletsprojects.com/ (дата обращения: 26.05.2026).',
    '[2] Jinja Documentation. URL: https://jinja.palletsprojects.com/ (дата обращения: 26.05.2026).',
    '[3] SQLAlchemy ORM Documentation. URL: https://docs.sqlalchemy.org/en/20/orm/ (дата обращения: 26.05.2026).',
    '[4] PostgreSQL Documentation. Data Definition. URL: https://www.postgresql.org/docs/current/ddl.html (дата обращения: 26.05.2026).',
    '[5] Docker Docs. Compose file reference. URL: https://docs.docker.com/reference/compose-file/ (дата обращения: 26.05.2026).',
    '[6] Werkzeug Documentation. Security Helpers. URL: https://werkzeug.palletsprojects.com/en/stable/utils/#module-werkzeug.security (дата обращения: 26.05.2026).',
    '[7] Requests Documentation. Quickstart. URL: https://requests.readthedocs.io/en/latest/user/quickstart/ (дата обращения: 26.05.2026).',
]
for src in sources:
    add_paragraph(doc, src)

doc.save(OUT)
print(OUT)
